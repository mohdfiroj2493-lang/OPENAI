from docx import Document
from .pdf_tools import extract_uploaded_pdf
from .text_utils import clean_text, extract_key_points, keyword_counter


def extract_docx(uploaded_file):
    doc = Document(uploaded_file)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join([cell.text.strip() for cell in row.cells]))
    return clean_text("\n".join(parts))


def extract_txt(uploaded_file):
    return uploaded_file.read().decode("utf-8", errors="ignore")


def extract_file_text(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_uploaded_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_docx(uploaded_file)
    return extract_txt(uploaded_file)


def document_report(text: str, focus: str = ""):
    query = focus or "important summary conclusions recommendations equations methods"
    points = extract_key_points(text, query, n=18)
    terms = keyword_counter(text, limit=15)
    out = []
    out.append("## Document Summary")
    out.append("### Key points")
    out.extend([f"- {p}" for p in points])
    out.append("\n### Important terms")
    out.append(", ".join([f"{w} ({c})" for w, c in terms]))
    return "\n".join(out)
